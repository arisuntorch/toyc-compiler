#!/usr/bin/env python3
"""Render the maintained Markdown practice report as a styled DOCX file."""

from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "practice_report_draft.md"
OUT = ROOT / "docs" / "practice_report.docx"


def set_east_asia_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name, size_pt=None, color=None, bold=None):
    font = style.font
    font.name = font_name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        font.size = Pt(size_pt)
    if color is not None:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold


def set_para_spacing(style, before=0, after=6, line=1.15):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
        set_cell_text(table.rows[0].cells[index], header, bold=True)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            set_cell_text(row.cells[index], value)
    doc.add_paragraph()


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    paragraph._p.get_or_add_pPr().append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_east_asia_font(run, "Consolas")
    run.font.size = Pt(9)


def markdown_cells(line):
    return [cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")]


def is_table_separator(line):
    cells = markdown_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def starts_block(lines, index):
    line = lines[index].strip()
    if not line:
        return True
    if line.startswith(("#", "```", "- ")) or re.match(r"\d+\.\s", line):
        return True
    return (
        line.startswith("|")
        and index + 1 < len(lines)
        and is_table_separator(lines[index + 1].strip())
    )


def join_wrapped_lines(parts):
    result = ""
    for raw in parts:
        part = raw.strip()
        if not part:
            continue
        if result:
            left = result[-1]
            right = part[0]
            left_word = left.isascii() and (left.isalnum() or left in "`_)")
            right_word = right.isascii() and (right.isalnum() or right in "`_(")
            if left_word or right_word:
                result += " "
        result += part
    return result


def render_markdown(doc, source):
    lines = source.splitlines()
    index = 0
    title_seen = False
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        if line.startswith("```"):
            index += 1
            block = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index].rstrip())
                index += 1
            index += 1
            add_code_block(doc, block)
            continue

        if (
            line.startswith("|")
            and index + 1 < len(lines)
            and is_table_separator(lines[index + 1].strip())
        ):
            headers = markdown_cells(line)
            index += 2
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = markdown_cells(lines[index])
                if len(cells) == len(headers):
                    rows.append(cells)
                index += 1
            add_table(doc, headers, rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not title_seen:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                set_east_asia_font(run, "黑体")
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string("0B2545")
                paragraph.paragraph_format.space_after = Pt(18)
                title_seen = True
            else:
                doc.add_heading(text, level=min(max(1, level - 1), 3))
            index += 1
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            parts = [line[2:]]
            index += 1
            while index < len(lines) and not starts_block(lines, index):
                parts.append(lines[index].strip())
                index += 1
            paragraph.add_run(join_wrapped_lines(parts).replace("`", ""))
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            parts = [numbered.group(1)]
            index += 1
            while index < len(lines) and not starts_block(lines, index):
                parts.append(lines[index].strip())
                index += 1
            paragraph.add_run(join_wrapped_lines(parts).replace("`", ""))
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines) and not starts_block(lines, index):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = doc.add_paragraph(join_wrapped_lines(paragraph_lines).replace("`", ""))
        paragraph.style = doc.styles["Normal"]


def add_page_number(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
    paragraph.add_run(" 页")


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    add_page_number(section)

    styles = doc.styles
    set_style_font(styles["Normal"], "宋体", 10.5, "000000")
    set_para_spacing(styles["Normal"], after=6, line=1.15)
    set_style_font(styles["Heading 1"], "黑体", 15, "2E74B5", True)
    set_para_spacing(styles["Heading 1"], before=14, after=7, line=1.1)
    set_style_font(styles["Heading 2"], "黑体", 12.5, "2E74B5", True)
    set_para_spacing(styles["Heading 2"], before=10, after=5, line=1.1)
    set_style_font(styles["Heading 3"], "黑体", 11.5, "1F4D78", True)
    set_para_spacing(styles["Heading 3"], before=8, after=4, line=1.1)
    set_style_font(styles["List Bullet"], "宋体", 10.5, "000000")
    set_para_spacing(styles["List Bullet"], after=3, line=1.1)
    set_style_font(styles["List Number"], "宋体", 10.5, "000000")
    set_para_spacing(styles["List Number"], after=3, line=1.1)

    render_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.core_properties.title = "编译系统实践报告：ToyC 到 RISC-V32 汇编编译器"
    doc.core_properties.subject = "ToyC compiler design, optimization, and verification"
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
